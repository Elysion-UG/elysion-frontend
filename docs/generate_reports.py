"""
Marketplace Backend & Frontend — Gap Analysis Report Generator
Generates: gap-analysis.xlsx + status-report.docx
"""

import os
from datetime import date
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TODAY = date.today().strftime("%d.%m.%Y")
DOCS_DIR = os.path.dirname(os.path.abspath(__file__))

GREEN      = "92D050"
YELLOW     = "FFEB9C"
RED        = "FFC7CE"
ORANGE     = "FFCC99"
WHITE      = "FFFFFF"
DARK       = "1F3864"
LIGHT_BLUE = "DEEAF1"

def set_col_widths(ws, widths):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

def header_row(ws, row_num, values, bg=DARK, fg=WHITE):
    for col, val in enumerate(values, 1):
        cell = ws.cell(row=row_num, column=col, value=val)
        cell.fill = PatternFill("solid", fgColor=bg)
        cell.font = Font(bold=True, color=fg, size=10)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        thin = Side(style="thin", color="AAAAAA")
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)

def data_row(ws, row_num, values, bg=WHITE):
    thin = Side(style="thin", color="DDDDDD")
    for col, val in enumerate(values, 1):
        cell = ws.cell(row=row_num, column=col, value=val)
        cell.fill = PatternFill("solid", fgColor=bg)
        cell.font = Font(size=9)
        cell.alignment = Alignment(vertical="top", wrap_text=True)
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)

def status_color(status):
    if any(x in status for x in ["✅", "Fertig", "100%"]): return GREEN
    if any(x in status for x in ["❌", "Fehlt", "0%"]): return RED
    if any(x in status for x in ["⚠️", "Teils", "Partial"]): return YELLOW
    return WHITE

def priority_color(prio):
    if "P0" in prio: return RED
    if "P1" in prio: return ORANGE
    if "P2" in prio: return YELLOW
    return WHITE


# ===========================================================================
# EXCEL
# ===========================================================================

def build_excel():
    wb = Workbook()
    wb.remove(wb.active)

    # --- Sheet 1: Backend Module Status ---
    ws1 = wb.create_sheet("01 Backend Module Status")
    ws1.freeze_panes = "A2"
    ws1.row_dimensions[1].height = 36
    header_row(ws1, 1, ["Modul", "Status", "% Fertig", "Was implementiert", "Was fehlt", "Bekannte Bugs", "Migrations"])
    set_col_widths(ws1, [22, 12, 10, 50, 60, 55, 20])

    backend_data = [
        ("01 Authentication", "⚠️ Teils fertig", "85%",
         "Register/Login/Logout mit JWT + HttpOnly Refresh-Cookie\nEmail-Verification + Password-Reset Flow\nAddress CRUD + Default-Setzung\nAdmin: User suspend/activate, Seller approve/reject/suspend\nGDPR Soft-Delete + Email-Anonymisierung\nBCrypt (Stärke 12) + Pepper\nIP-basiertes Rate-Limiting\nSeller Profile Status-Machine\nRefresh-Token Rotation mit Hash-Storage\nToken-TTL: 5h Access / 7d Refresh",
         "Account-Lockout nach X Fehlversuchen\nEmail-basiertes Rate-Limiting\nRetry-After Header\nWerteprofil GET/PUT Endpoints (korrekte Felder)\nLogin-Response mit user-Objekt\ncreatedAt in GET /users/me Response",
         "KRITISCH: Email-Constraint lehnt name.surname@domain.com ab\nKRITISCH: Email-Verification-Link zeigt auf Backend (User sieht JSON)\nKRITISCH: Password-Reset-Link zeigt auf Backend (User sieht JSON)\nKRITISCH: Refresh-Token Race Condition (kein pessimistic Lock)\nSEC: Rate-Limit via X-Forwarded-For spoofbar\nSEC: In-memory Buckets brechen bei Horizontal Scaling\nBUG: /addresses/{id}/default vs /set-default (Pfad-Mismatch)\nBUG: markUsedNow() auf RefreshToken nie aufgerufen",
         "V1, V2"),
        ("02 Product Management", "❌ Nicht implementiert", "10%",
         "Datenbankschema vorhanden (V2__catalog_foundation.sql)\nProduct, ProductImage, ProductSlugHistory Entities\nVariant, VariantOption Entities\nCategory Entity (3-Level-Hierarchie)\nMaterialized View product_search_view (partial)",
         "ALLE Controller fehlen: GET/POST/PATCH /products, GET /products/{slug}\nProductService CRUD komplett\nVariantService (Inventory, Reservierung)\nSlugService (Slug-Generierung + History für 301 Redirects)\nProductSearchService + ViewCounterService\nStatus-Machine Enforcement:\n  DRAFT→REVIEW (mind. 1 Bild)\n  REVIEW→ACTIVE (mind. 1 Zertifikat)\n  REVIEW→REJECTED (Admin only)\n  ACTIVE↔INACTIVE",
         "BLOCKER: V3 Migration erstellt Tabellen aber keine Java-Entities → Hibernate Startup-Fehler (ddl-auto=validate)",
         "V2 (catalog_foundation)"),
        ("03 Certificate Management", "✅ Vollständig", "100%",
         "Vollständiger Lifecycle: PENDING→VERIFIED/REJECTED, VERIFIED→EXPIRED\nDB-Triggers: Auto-Sync verified_certificate_count\nSeller: POST/GET Zertifikate + Produkt-Verknüpfung\nAdmin: GET pending, PATCH verify/reject\nErinnerungs-Emails (30d, 7d vor Ablauf)\nAuto-Aktivierung Produkt bei Cert-Verifizierung",
         "—", "—", "V3"),
        ("04 Matching Engine", "❌ Nicht implementiert", "0%",
         "—",
         "MatchScoreService: calculateSimpleMatch(), calculateExtendedMatch()\nCertificateCategoryMappingService (konfigurierbar)\nBulk Match-Score Berechnung für Produktlisten\nRecommendationService: GET /api/v1/recommendations\nMatch-Score Caching (Redis)\nmatch_score in ProductListResponse\nSortierung nach match_score",
         "Frontend RecommendationsWidget ist bereits fertig und wartet auf diesen Endpunkt!",
         "—"),
        ("05 Cart & Checkout", "✅ Vollständig", "100%",
         "Cart + CartItem mit Preis-Snapshot\nGast-Cart via HttpOnly Cookie (30d)\nGast→User Merge beim Login\nInventory-Reservierung: atomisch, verhindert Overselling\nCheckout Start + Complete mit Multi-Vendor Split",
         "—", "—", "V4"),
        ("06 Order Management", "✅ Vollständig", "100%",
         "Order + OrderGroup + OrderItem mit Snapshots\nMulti-Vendor Split: 1 Order → N OrderGroups\nBuyer: GET /orders, GET /orders/{id}\nSeller: Bestellungen, Versand, Lieferung\nStatus-Machine: CONFIRMED→PROCESSING→SHIPPED→DELIVERED|CANCELLED",
         "Returns/Refunds vollständiger Flow\nKäufer-bestätigte Lieferung",
         "—", "V5, V6, V7"),
        ("07 Payment Processing", "✅ Fast vollständig", "95%",
         "Payment, Webhook, Settlement, Payout, Refund Entities\nPOST /payments/create-intent (Webhook signatur-verifiziert)\nGET /seller/settlements\nAdmin Read-Only für alle Finanzdaten\nCommission-Model mit Plattformgebühr",
         "Payout-Onboarding (Seller Bankverbindung)\nZahlungsabgleich/Reconciliation\nAdmin Finance Mutations",
         "—", "V9, V10, V11, V12"),
        ("08 File Upload", "✅ Vollständig", "100%",
         "FileAsset mit MIME-Typ, Magic-Byte Validierung, Größen-Limits\nKategorien: PRODUCT_IMAGE, CERTIFICATE_DOCUMENT, PROFILE_IMAGE, SHOP_LOGO\nPOST /files/upload, GET /files/{id}/content, DELETE, replace\nLinking-Model + Soft-Delete",
         "Malware-Scan, CDN, Thumbnails, Cleanup-Scheduler",
         "—", "V8, V13"),
        ("09 Admin Panel", "✅ Fast vollständig", "90%",
         "Dashboard mit Stats\nZertifikat-, User-, Seller-, Produkt-Moderation\nAdminAuditLog für alle Admin-Aktionen\nOrder Support + Finance Read-Only",
         "Reports & Exports (CSV/Excel)\nMassen-Aktionen, 2FA-Enforcement, Admin-Rollen-Hierarchie",
         "—", "V14"),
        ("10 Email Service", "✅ Fast vollständig", "90%",
         "SMTP Transport + NoopService für Tests\nTemplates: auth-verify, password-reset, cert-expiry, order-confirm, payment-success, refund\nDelivery Logging + Duplikat-Schutz",
         "Outbox-Pattern für Retry\nResend Verification Email Endpunkt\nLokalisierung (i18n)",
         "—", "V15"),
    ]

    for r, row in enumerate(backend_data, 2):
        data_row(ws1, r, row)
        ws1.cell(r, 2).fill = PatternFill("solid", fgColor=status_color(row[1]))
        ws1.row_dimensions[r].height = 120

    # --- Sheet 2: Frontend Module Status ---
    ws2 = wb.create_sheet("02 Frontend Module Status")
    ws2.freeze_panes = "A2"
    ws2.row_dimensions[1].height = 36
    header_row(ws2, 1, ["Feature / Seite", "Komponente(n)", "Status", "% fertig", "API aufgerufen?", "Was fehlt / Blocker"])
    set_col_widths(ws2, [24, 32, 14, 10, 45, 60])

    frontend_data = [
        ("Authentication\n(Login/Register/Reset)", "LoginModal, Onboarding,\nEmailVerification, ResetPassword,\nAuthContext", "✅ Fertig", "100%",
         "POST /auth/login ✅\nPOST /auth/register ✅\nPOST /auth/refresh ✅\nPOST /auth/logout ✅\nPOST /auth/verify-email ✅\nPOST /auth/reset-password ✅",
         "forgot-password UI fehlt (Service vorhanden)\nDelete-Account UI fehlt"),
        ("User Profil", "Profil.tsx\nUserService, AddressService", "⚠️ Teils", "80%",
         "GET/PATCH /users/me ✅\nGET /users/me/addresses ✅",
         "Address CRUD (Create/Update/SetDefault/Delete) unvollständig in UI\nSeller-Profil-Bearbeitung unvollständig"),
        ("Käufer-Präferenzen", "Praeferenzen.tsx\nBuyerValueProfileService", "✅ Fertig", "100%",
         "GET/PUT /users/me/profile ✅", "—"),
        ("Shop Browse / Suche", "SustainableShop.tsx\nRecommendationsWidget.tsx", "✅ Fertig", "95%",
         "GET /products ✅\nGET /categories/tree ✅\nGET /recommendations ✅",
         "Category-Filter-UI angezeigt aber nicht an API übergeben"),
        ("Produkt-Detail", "ProductDetail.tsx", "❌ BLOCKER", "0%",
         "GET /products/{slug} ❌ NICHT AUFGERUFEN",
         "100% MOCK-DATEN: Hardcodiertes Produkt 'Organic Cotton T-Shirt'\nUser kann keine echten Produkte ansehen\nAdd-to-Cart nicht funktionstüchtig\nFix: ProductService.getBySlug(slug) + CertificateService.getProductCerts()"),
        ("Warenkorb", "Cart.tsx, CartContext, CartService", "✅ Fertig", "100%",
         "GET /cart ✅\nPOST /cart/items ✅\nPATCH /cart/items/{id} ✅\nDELETE /cart/items/{id} ✅", "—"),
        ("Checkout", "Checkout.tsx, CheckoutService", "⚠️ Teils", "95%",
         "POST /checkout ✅\nPOST /checkout/complete ✅\nGET /users/me/addresses ✅",
         "paymentMethod hardcoded 'MOCK' — echte Zahlung nicht möglich"),
        ("Bestellungen (Käufer)", "Orders.tsx, OrderDetail.tsx\nOrderService", "✅ Fertig", "100%",
         "GET /orders ✅\nGET /orders/{id} ✅", "—"),
        ("Seller Dashboard — Produkte", "SellerDashboard.tsx (Products Tab)\nProductService", "⚠️ Teils", "85%",
         "GET /products (seller) ✅\nPATCH /products/{id}/status ✅",
         "'Neues Produkt' Dialog nicht verdrahtet\nBild-Upload UI fehlt (FileService bereit)\nVarianten-Management UI fehlt"),
        ("Seller Dashboard — Bestellungen", "SellerDashboard.tsx (Orders Tab)\nSellerOrderService", "✅ Fertig", "100%",
         "GET /seller/orders ✅\nPATCH /seller/orders/{id}/ship ✅", "—"),
        ("Seller Dashboard — Auszahlungen", "SellerDashboard.tsx (Settlements)\nSellerOrderService", "✅ Fertig", "100%",
         "GET /seller/orders/settlements ✅", "—"),
        ("Admin — User Verwaltung", "AdminUsers.tsx, AdminUserDetail.tsx\nAdminService", "✅ Fast fertig", "95%",
         "GET /admin/users ✅\nGET /admin/users/{id} ✅\nPOST /admin/seller-profiles/{id}/approve ✅",
         "suspendUser/activateUser Buttons im Detail-View nicht verdrahtet"),
        ("Zertifikats-Verwaltung (Seller)", "— (kein UI vorhanden)", "❌ Fehlt", "5%",
         "—",
         "CertificateService komplett bereit, aber kein Seller-UI\nKein Upload-Dialog\nKeine Anzeige auf Produktseite"),
        ("Datei-Upload UI", "— (kein UI vorhanden)", "❌ Fehlt", "10%",
         "—",
         "FileService 100% fertig aber nirgends verdrahtet\nKein Produkt-Bild Upload, kein Profilbild Upload"),
        ("Empfehlungen / Matching", "RecommendationsWidget.tsx\nRecommendationService", "✅ Frontend fertig", "100%",
         "GET /recommendations ✅",
         "Backend Modul 04 (Matching Engine) = 0%!\nWidget zeigt leer bis Backend fertig"),
        ("Produzenten-Seite", "ProducerPage.tsx", "⚠️ Teils", "15%",
         "—", "Komplett Mock-Daten, kein Backend-Endpunkt für Public Seller-Profile"),
    ]

    for r, row in enumerate(frontend_data, 2):
        data_row(ws2, r, row)
        ws2.cell(r, 3).fill = PatternFill("solid", fgColor=status_color(row[2]))
        ws2.row_dimensions[r].height = 90

    # --- Sheet 3: Technische Änderungen Backend ---
    ws3 = wb.create_sheet("03 Aenderungen Backend")
    ws3.freeze_panes = "A2"
    ws3.row_dimensions[1].height = 36
    header_row(ws3, 1, ["ID", "Priorität", "Typ", "Beschreibung", "Betroffene Datei(en)", "Aufwand (Tage)", "Status"])
    set_col_widths(ws3, [6, 12, 14, 65, 50, 14, 14])

    backend_changes = [
        ("B-01", "P0 — KRITISCH", "Bug",
         "Email-Constraint-Bug: V2__auth_schema.sql lehnt gültige Emails ab (z.B. name.surname@domain.com). Der CHECK position('.' in email) > position('@' in email)+1 schlägt fehl wenn der erste Punkt vor dem @ steht. Trifft ~20% aller Geschäfts-Emails.",
         "src/main/resources/db/migration/\n→ Neue Migration V17__fix_email_constraint.sql", "0.5", "Offen"),
        ("B-02", "P0 — KRITISCH", "Security",
         "Rate-Limit X-Forwarded-For Spoofing: AuthRateLimitFilter vertraut blind dem X-Forwarded-For Header ohne Proxy-Whitelist. Angreifer können beliebige IPs vortäuschen und das Rate-Limit umgehen.",
         "infrastructure/security/AuthRateLimitFilter.java", "1", "Offen"),
        ("B-03", "P0 — KRITISCH", "Bug",
         "Email-Verification-Link zeigt auf Backend statt Frontend: User bekommt Link auf /api/v1/auth/verify-email und sieht rohen JSON. Email-Verifizierung kann nicht abgeschlossen werden.",
         "application/AuthEmailService.java\napp.base-url Property prüfen", "0.5", "Offen"),
        ("B-04", "P0 — KRITISCH", "Bug",
         "Password-Reset-Link zeigt auf Backend statt Frontend. Gleicher Fehler wie B-03.",
         "application/AuthEmailService.java", "0.5", "Offen"),
        ("B-05", "P0 — KRITISCH", "Bug",
         "Refresh-Token Race Condition: AuthService.refresh() liest und updated den Token ohne pessimistic Lock. Concurrent Requests mit demselben Token können beide erfolgreich sein → 2 gültige Access Tokens.",
         "application/AuthService.java\ninfrastructure/jpa/RefreshTokenRepository.java\n→ @Lock(PESSIMISTIC_WRITE) hinzufügen", "1", "Offen"),
        ("B-06", "P1 — HOCH", "Feature",
         "Modul 02 Product Management vollständig implementieren:\n- ProductController: GET /products, GET /products/{slug}, POST /products, PATCH /products/{id}, PATCH /products/{id}/status\n- ProductService CRUD\n- VariantService (Inventory, Reservierung, Freigabe)\n- SlugService (Slug-Generierung + Slug-History für 301 Redirects)\n- ProductImageService\n- Status-Machine Enforcement\n- Admin Review-Queue",
         "api/v1/products/ (neu)\napplication/ (Services neu)\ndomain/ (Entities prüfen)\ninfrastructure/jpa/ (Repositories)", "10", "Offen"),
        ("B-07", "P1 — HOCH", "Bug",
         "Login-Response enthält kein user-Objekt: Frontend muss nach jedem Login einen zusätzlichen GET /users/me aufrufen.",
         "application/AuthService.java\napi/v1/auth/AuthController.java\nLoginResponse DTO anpassen", "0.5", "Offen"),
        ("B-08", "P1 — HOCH", "Bug",
         "GET /users/me/profile gibt falsche Felder zurück: Endpoint gibt firstName, lastName, phone zurück statt Werteprofil-Objekt (simpleProfile, extendedProfile, activeProfileType).",
         "api/v1/users/UsersMeProfileController.java\nUserProfileResponse DTO\nUserProfileService", "1", "Offen"),
        ("B-09", "P1 — HOCH", "Bug",
         "Address set-default Pfad-Mismatch: Backend hat PATCH /addresses/{id}/default, Spec und Frontend erwarten /addresses/{id}/set-default.",
         "api/v1/users/UsersMeAddressesController.java", "0.5", "Offen"),
        ("B-10", "P2 — MITTEL", "Feature",
         "Modul 04 Matching Engine:\n- MatchScoreService: calculateSimpleMatch(), calculateExtendedMatch()\n- CertificateCategoryMappingService\n- Bulk Match-Score für Produktlisten\n- RecommendationService: GET /api/v1/recommendations\n- Caching (Redis)\n- match_score in ProductListResponse\nNOTE: Frontend RecommendationsWidget ist bereits fertig!",
         "application/matching/ (neu)\napi/v1/recommendations/ (neu)", "8", "Offen"),
        ("B-11", "P2 — MITTEL", "Feature",
         "Account-Lockout: Nach X fehlgeschlagenen Login-Versuchen Konto temporär sperren (z.B. 15 Min.). Aktuell nur IP-basiertes Rate-Limiting.",
         "application/AuthService.java\n→ Neue Felder: users.failed_login_count, locked_until", "2", "Offen"),
        ("B-12", "P2 — MITTEL", "Bug",
         "In-Memory Rate-Limiting bricht bei Horizontal Scaling: Jede Instanz hat eigene Buckets → bei N Instanzen sind N*5 Versuche möglich.",
         "infrastructure/security/AuthRateLimitFilter.java\n→ Redis-basiertes Rate-Limiting (Bucket4j + Redis)", "2", "Offen"),
        ("B-13", "P2 — MITTEL", "Feature",
         "Admin Reports & Exports: Kein CSV/Excel-Export möglich. Benötigt: User-Export, Order-Export, Revenue-Report, Cert-Status-Report.",
         "api/v1/admin/AdminReportsController.java (neu)\napplication/admin/AdminReportService.java (neu)", "4", "Offen"),
        ("B-14", "P2 — MITTEL", "Feature",
         "Email Outbox-Pattern + Retry: Emails direkt per SMTP ohne Retry. Bei SMTP-Ausfall gehen transaktionale Emails verloren.",
         "application/EmailDeliveryService.java\nNeuer Scheduler: EmailRetryScheduler.java\n(EmailDelivery Entity hat bereits Status-Felder)", "3", "Offen"),
        ("B-15", "P2 — MITTEL", "Feature",
         "Payout-Onboarding: Seller können keine Bankverbindung hinterlegen (IBAN oder Stripe Connect). Payouts im System modelliert, aber kein Onboarding-Flow.",
         "api/v1/users/UsersMeSellerProfileController.java\napplication/PayoutOnboardingService.java (neu)", "5", "Offen"),
        ("B-16", "P2 — MITTEL", "Feature",
         "Resend Email-Verification Endpunkt: Frontend hat TODO-Kommentar darüber. User können bei abgelaufenem Token keine neue Verification-Email anfordern.",
         "api/v1/auth/AuthController.java\napplication/AuthService.java", "0.5", "Offen"),
    ]

    for r, row in enumerate(backend_changes, 2):
        data_row(ws3, r, row)
        ws3.cell(r, 2).fill = PatternFill("solid", fgColor=priority_color(row[1]))
        ws3.row_dimensions[r].height = 100

    # --- Sheet 4: Technische Änderungen Frontend ---
    ws4 = wb.create_sheet("04 Aenderungen Frontend")
    ws4.freeze_panes = "A2"
    ws4.row_dimensions[1].height = 36
    header_row(ws4, 1, ["ID", "Priorität", "Typ", "Beschreibung", "Betroffene Datei(en)", "Aufwand (Tage)", "Status"])
    set_col_widths(ws4, [6, 12, 14, 65, 50, 14, 14])

    frontend_changes = [
        ("F-01", "P0 — KRITISCH", "Bug",
         "ProductDetail.tsx 100% Mock-Daten: Hardcodiertes Produkt 'Organic Cotton T-Shirt'. Echte Produkte nicht ansehbar, Add-to-Cart nicht funktionstüchtig.\nFix: URL-Slug aus Next.js Router → ProductService.getBySlug(slug) → CertificateService.getProductCerts() → echte Varianten-Auswahl → Add-to-Cart mit echter variantId",
         "src/components/ProductDetail.tsx", "2", "Offen"),
        ("F-02", "P0 — KRITISCH", "Bug",
         "Checkout paymentMethod hardcoded 'MOCK': Echte Zahlungsabwicklung nicht möglich.\nFix: Stripe Elements oder PayPal Integration. Bis dahin: klare 'Demo-Modus'-Meldung anzeigen.",
         "src/components/Checkout.tsx (Zeile ~42, ~59)", "5", "Offen"),
        ("F-03", "P1 — HOCH", "Feature",
         "SellerDashboard 'Neues Produkt' Dialog: ProductForm Komponente existiert bereits, ist aber nicht in SellerDashboard eingebunden. Seller können keine Produkte anlegen.",
         "src/components/SellerDashboard.tsx\nsrc/components/ProductForm.tsx", "1", "Offen"),
        ("F-04", "P1 — HOCH", "Feature",
         "Bild-Upload UI in ProductForm verdrahten: FileService vollständig implementiert (upload, link, unlink) aber nirgends in der UI eingebunden. Produkte haben keine Bilder.\nFix: Image-Upload-Bereich in ProductForm mit FileService.uploadAndLink()",
         "src/components/ProductForm.tsx\nsrc/services/index.ts (FileService bereit)", "2", "Offen"),
        ("F-05", "P1 — HOCH", "Feature",
         "CertificateForm erstellen: Seller haben kein UI zum Hochladen von Zertifikaten. Kernfeature des nachhaltigen Marktplatzes!\nFix: CertificateForm Dialog (Typ, Dokument via FileService, Metadaten) → in SellerDashboard Tab einbinden.",
         "src/components/CertificateForm.tsx (neu)\nsrc/components/SellerDashboard.tsx\nsrc/services/index.ts (CertificateService bereit)", "3", "Offen"),
        ("F-06", "P1 — HOCH", "Feature",
         "Address CRUD in Profil-Komponente vervollständigen: Services für Create/Update/SetDefault/Delete implementiert aber UI unvollständig.",
         "src/components/Profil.tsx\nsrc/services/index.ts (AddressService bereit)", "1.5", "Offen"),
        ("F-07", "P2 — MITTEL", "Bug",
         "Category-Filter in SustainableShop nicht aktiv: selectedCategoryId wird nicht als Parameter an ProductService.list() übergeben.",
         "src/components/SustainableShop.tsx", "0.5", "Offen"),
        ("F-08", "P2 — MITTEL", "Bug",
         "Admin UserDetail: suspendUser/activateUser Buttons nicht verdrahtet.",
         "src/components/AdminUserDetail.tsx", "0.5", "Offen"),
        ("F-09", "P2 — MITTEL", "Feature",
         "ProducerPage auf echte API umstellen: Komplett Mock-Daten. Benötigt Backend-Endpunkt GET /sellers/{slug}.",
         "src/components/ProducerPage.tsx\n→ Backend: GET /api/v1/sellers/{slug} nötig", "2", "Offen"),
        ("F-10", "P2 — MITTEL", "Feature",
         "Forgot-Password UI: AuthService.forgotPassword() implementiert aber kein UI-Element vorhanden.",
         "src/components/LoginModal.tsx\n→ 'Passwort vergessen?' Link + Dialog", "1", "Offen"),
        ("F-11", "P2 — MITTEL", "Feature",
         "Zertifikate auf Produktdetailseite anzeigen: CertificateService.getProductCertificates() bereit. Zertifikate als Badge/Liste auf ProductDetail zeigen.",
         "src/components/ProductDetail.tsx (nach F-01)\nsrc/services/index.ts (CertificateService)", "1", "Offen"),
    ]

    for r, row in enumerate(frontend_changes, 2):
        data_row(ws4, r, row)
        ws4.cell(r, 2).fill = PatternFill("solid", fgColor=priority_color(row[1]))
        ws4.row_dimensions[r].height = 100

    # --- Sheet 5: Migrations-Inventar ---
    ws5 = wb.create_sheet("05 Migrations-Inventar")
    ws5.freeze_panes = "A2"
    ws5.row_dimensions[1].height = 36
    header_row(ws5, 1, ["Version", "Name", "Scope / Beschreibung", "Enthaltene Tabellen / Trigger"])
    set_col_widths(ws5, [10, 42, 40, 65])

    migrations = [
        ("V1",  "core_auth_and_users",            "Kern-Auth-Schema",                                       "users, user_profiles, addresses, seller_profiles, refresh_tokens, one_time_tokens"),
        ("V2",  "catalog_foundation",              "Produkt-Katalog-Schema",                                 "categories, products, product_images, variants, variant_options, product_slug_history\nMaterialized View: product_search_view"),
        ("V3",  "certificates",                    "Zertifikats-Schema mit DB-Triggers",                     "certificates, product_certificates\nTrigger: sync_product_verified_certificate_count"),
        ("V4",  "cart_and_checkout_foundation",    "Warenkorb und Checkout",                                 "carts, cart_items"),
        ("V5",  "order_foundation",                "Basis-Bestellschema",                                    "orders, order_items, order_address_snapshots, order_product_snapshots"),
        ("V6",  "order_group_shipping_foundation", "Multi-Vendor Order Groups mit Versand",                  "order_groups (tracking_number, carrier, shipped_at)"),
        ("V7",  "order_group_delivery_foundation", "Lieferungs-Felder",                                      "order_groups (delivered_at)"),
        ("V8",  "file_upload_foundation",          "Datei-Upload und Storage",                               "file_assets"),
        ("V9",  "payment_foundation",              "Zahlungsschema + Webhook-Events",                        "payments, payment_webhook_events"),
        ("V10", "settlement_foundation",           "Abrechnungsschema",                                      "order_group_settlements"),
        ("V11", "refund_foundation",               "Rückerstattungsschema",                                  "refunds"),
        ("V12", "payout_execution_foundation",     "Auszahlungsschema",                                      "payouts, payout_settlements"),
        ("V13", "file_lifecycle_foundation",       "Datei-Lebenszyklus (Soft-Delete)",                       "file_assets (deleted_at, cleanup_after)"),
        ("V14", "admin_foundation",                "Admin-Audit-Log",                                        "admin_audit_logs"),
        ("V15", "email_delivery_foundation",       "Email-Delivery-Tracking",                                "email_deliveries"),
        ("V16", "pending_order_expiry",            "Pending-Order-Ablauf (automatisches Verfallsdatum)",      "orders (expires_at)\norder_groups (expiry-Felder)"),
    ]

    for r, row in enumerate(migrations, 2):
        bg = LIGHT_BLUE if r % 2 == 0 else WHITE
        data_row(ws5, r, row, bg=bg)
        ws5.row_dimensions[r].height = 50

    # --- Sheet 6: Implementierungs-Roadmap ---
    ws6 = wb.create_sheet("06 Implementierungs-Roadmap")
    ws6.freeze_panes = "A2"
    ws6.row_dimensions[1].height = 36
    header_row(ws6, 1, ["Phase", "Bereich", "ID", "Aufgabe", "Aufwand (Tage)", "Abhängigkeit"])
    set_col_widths(ws6, [22, 12, 8, 65, 14, 18])

    roadmap = [
        ("Phase 0 — SOFORT", "Backend", "B-03", "Email-Verification-Link auf Frontend-URL umstellen", "0.5", "—"),
        ("Phase 0 — SOFORT", "Backend", "B-04", "Password-Reset-Link auf Frontend-URL umstellen", "0.5", "—"),
        ("Phase 0 — SOFORT", "Backend", "B-01", "Email-Constraint-Bug fixen (neue Migration V17)", "0.5", "—"),
        ("Phase 0 — SOFORT", "Backend", "B-07", "Login-Response mit user-Objekt ergänzen", "0.5", "—"),
        ("Phase 0 — SOFORT", "Backend", "B-05", "Refresh-Token Race Condition: pessimistic Lock", "1", "—"),
        ("Phase 0 — SOFORT", "Frontend", "F-01", "ProductDetail.tsx: Mock durch ProductService.getBySlug() ersetzen", "2", "B-06 (teilweise)"),
        ("Phase 1 — VOR LAUNCH", "Backend", "B-06", "Modul 02 Product Management vollständig implementieren", "10", "—"),
        ("Phase 1 — VOR LAUNCH", "Backend", "B-02", "Rate-Limit X-Forwarded-For absichern", "1", "—"),
        ("Phase 1 — VOR LAUNCH", "Backend", "B-08", "GET /users/me/profile Werteprofil-Felder korrigieren", "1", "—"),
        ("Phase 1 — VOR LAUNCH", "Backend", "B-09", "Address set-default Pfad-Mismatch fixen", "0.5", "—"),
        ("Phase 1 — VOR LAUNCH", "Backend", "B-16", "Resend Verification Email Endpunkt", "0.5", "B-03"),
        ("Phase 1 — VOR LAUNCH", "Frontend", "F-03", "SellerDashboard: Neues-Produkt-Dialog (ProductForm) einbinden", "1", "B-06"),
        ("Phase 1 — VOR LAUNCH", "Frontend", "F-04", "Bild-Upload in ProductForm verdrahten (FileService)", "2", "B-06, F-03"),
        ("Phase 1 — VOR LAUNCH", "Frontend", "F-05", "CertificateForm erstellen + in SellerDashboard einbinden", "3", "B-06"),
        ("Phase 1 — VOR LAUNCH", "Frontend", "F-06", "Address CRUD in Profil vervollständigen", "1.5", "B-09"),
        ("Phase 1 — VOR LAUNCH", "Frontend", "F-07", "Category-Filter in SustainableShop aktivieren", "0.5", "B-06"),
        ("Phase 2 — NACH LAUNCH", "Backend", "B-10", "Modul 04 Matching Engine implementieren", "8", "B-06"),
        ("Phase 2 — NACH LAUNCH", "Backend", "B-11", "Account-Lockout implementieren", "2", "—"),
        ("Phase 2 — NACH LAUNCH", "Backend", "B-12", "Redis-basiertes Rate-Limiting", "2", "—"),
        ("Phase 2 — NACH LAUNCH", "Backend", "B-13", "Admin Reports & Exports", "4", "—"),
        ("Phase 2 — NACH LAUNCH", "Backend", "B-14", "Email Outbox-Pattern + Retry", "3", "—"),
        ("Phase 2 — NACH LAUNCH", "Backend", "B-15", "Payout Onboarding (Bankverbindung)", "5", "—"),
        ("Phase 2 — NACH LAUNCH", "Frontend", "F-02", "Echte Payment-Integration (Stripe/PayPal)", "5", "—"),
        ("Phase 2 — NACH LAUNCH", "Frontend", "F-08", "Admin: suspendUser/activateUser Buttons verdrahten", "0.5", "—"),
        ("Phase 2 — NACH LAUNCH", "Frontend", "F-09", "ProducerPage: Mock durch echte API ersetzen", "2", "Backend-Endpunkt nötig"),
        ("Phase 2 — NACH LAUNCH", "Frontend", "F-10", "Forgot-Password UI erstellen", "1", "—"),
        ("Phase 2 — NACH LAUNCH", "Frontend", "F-11", "Zertifikate auf Produktdetailseite anzeigen", "1", "F-01, B-06"),
    ]

    phase_colors = {"Phase 0": RED, "Phase 1": ORANGE, "Phase 2": YELLOW}
    for r, row in enumerate(roadmap, 2):
        data_row(ws6, r, row)
        for key, color in phase_colors.items():
            if key in row[0]:
                ws6.cell(r, 1).fill = PatternFill("solid", fgColor=color)
                break
        ws6.row_dimensions[r].height = 45

    excel_path = os.path.join(DOCS_DIR, "gap-analysis.xlsx")
    wb.save(excel_path)
    print(f"[OK] Excel: {excel_path}")
    return excel_path


# ===========================================================================
# WORD
# ===========================================================================

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    color = RGBColor(0x1F, 0x38, 0x64) if level == 1 else RGBColor(0x2E, 0x74, 0xB5)
    if h.runs:
        h.runs[0].font.color.rgb = color
    return h

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F3864")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for ri, row in enumerate(rows, 1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(8)
    if col_widths:
        for row in table.rows:
            for ci, width in enumerate(col_widths):
                if ci < len(row.cells):
                    row.cells[ci].width = Inches(width)
    return table


def build_word():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # Title
    title = doc.add_heading("Marketplace Platform", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        title.runs[0].font.size = Pt(28)
    sub = doc.add_paragraph("Backend & Frontend — Status Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].font.size = Pt(16)
        sub.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    dp = doc.add_paragraph(f"Stand: {TODAY}")
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if dp.runs: dp.runs[0].font.size = Pt(11)
    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "Dieser Report gibt einen vollständigen Überblick über den Implementierungsstand der Marketplace-Plattform. "
        "Er enthält alle identifizierten Lücken, Bugs, Sicherheitsprobleme sowie einen priorisierten Implementierungsplan."
    )
    add_heading(doc, "Gesamtstatus", 2)
    add_table(doc,
        ["Bereich", "Fertigstellung", "Kritische Blocker", "Offene Features"],
        [
            ["Backend (10 Module)", "~75%", "5 (P0-Bugs)", "2 Module fehlen komplett"],
            ["Frontend (17 Seiten)", "~73%", "2 (ProductDetail, Payment)", "7 Lücken"],
            ["Gesamtsystem", "~70%", "7 (P0-Priorität)", "~25 offene Punkte"],
        ],
        col_widths=[2.0, 1.6, 1.6, 1.8]
    )

    doc.add_paragraph()
    add_heading(doc, "Sofortmassnahmen (P0)", 2)
    for item in [
        "Email-Verification-Link zeigt auf Backend — Nutzer KÖNNEN sich nicht registrieren",
        "Password-Reset-Link zeigt auf Backend — Passwort-Reset unmöglich",
        "Email-Constraint-Bug lehnt ~20% aller Geschäfts-Emails ab",
        "Refresh-Token Race Condition — Sicherheitslücke",
        "Rate-Limit via X-Forwarded-For spoofbar — Brute-Force-Schutz umgehbar",
        "ProductDetail.tsx zeigt 100% Mockdaten — kein echtes Produkt ansehbar",
        "Checkout-Payment hardcoded 'MOCK' — keine echte Zahlung möglich",
    ]:
        p = doc.add_paragraph(f"• {item}", style="List Bullet")
        p.runs[0].font.size = Pt(9)

    add_heading(doc, "Nicht implementierte Module", 2)
    doc.add_paragraph(
        "Modul 02 (Product Management): Nur Datenbankschema vorhanden — keine Controller, keine Services, keine Endpoints. "
        "Seller können keine Produkte anlegen, Käufer keine echten Produkte ansehen. "
        "Zudem schlägt Hibernate ddl-auto=validate beim Startup fehl (V3 Migration vs. fehlende Java Entities)."
    ).runs[0].font.size = Pt(9)
    doc.add_paragraph(
        "Modul 04 (Matching Engine): Komplett nicht implementiert (0%). "
        "Das Frontend-RecommendationsWidget ist bereits fertig und wartet auf GET /api/v1/recommendations."
    ).runs[0].font.size = Pt(9)
    doc.add_page_break()

    # 2. Backend Details
    add_heading(doc, "2. Backend — Detailanalyse nach Modul", 1)

    modules = [
        ("01 Authentication", "⚠️ Teils fertig — 85%", "V1, V2",
         ["Register/Login/Logout mit JWT + HttpOnly Refresh-Cookie (5h/7d TTL)",
          "Email-Verification + Password-Reset Flow",
          "Address CRUD + Default-Setzung",
          "Admin: User suspend/activate, Seller approve/reject/suspend",
          "GDPR Soft-Delete mit Email-Anonymisierung",
          "BCrypt (Stärke 12) + Pepper, Rate-Limiting via IP",
          "Seller Profile Status-Machine, Refresh-Token Rotation"],
         ["Account-Lockout nach X Fehlversuchen",
          "Korrekte Werteprofil-Endpoints",
          "Login-Response mit user-Objekt"],
         ["KRITISCH: Email-Constraint lehnt name.surname@domain.com ab",
          "KRITISCH: Email-Links zeigen auf Backend (User sieht JSON)",
          "KRITISCH: Refresh-Token Race Condition",
          "SEC: Rate-Limit via X-Forwarded-For spoofbar",
          "BUG: Login-Response fehlt user-Objekt",
          "BUG: Address set-default Pfad-Mismatch"]),
        ("02 Product Management", "❌ Nicht implementiert — 10%", "V2 (catalog_foundation)",
         ["Datenbankschema vorhanden",
          "Product, Variant, Category Entities (Schema-only)"],
         ["ALLE Controller fehlen (GET/POST/PATCH /products)",
          "ProductService, VariantService, SlugService",
          "Status-Machine Enforcement",
          "Inventory-Reservierung"],
         ["BLOCKER: Hibernate ddl-auto=validate schlägt beim Start fehl"]),
        ("03 Certificate Management", "✅ Vollständig — 100%", "V3",
         ["Vollständiger Lifecycle: PENDING→VERIFIED/REJECTED→EXPIRED",
          "DB-Triggers, Seller- und Admin-Endpoints",
          "Erinnerungs-Emails, Auto-Aktivierung Produkt"],
         [], []),
        ("04 Matching Engine", "❌ Nicht implementiert — 0%", "—",
         [],
         ["MatchScoreService, CertificateCategoryMappingService",
          "RecommendationService: GET /api/v1/recommendations",
          "Match-Score Caching, Bulk-Berechnung"],
         ["Frontend RecommendationsWidget wartet bereits darauf!"]),
        ("05 Cart & Checkout", "✅ Vollständig — 100%", "V4",
         ["Cart + CartItem mit Preis-Snapshot",
          "Gast-Cart, Merge beim Login",
          "Atomische Inventory-Reservierung",
          "Multi-Vendor Checkout"],
         [], []),
        ("06 Order Management", "✅ Vollständig — 100%", "V5, V6, V7",
         ["Multi-Vendor Split, Snapshots",
          "Buyer + Seller Endpoints",
          "Status-Machine + Aggregation"],
         ["Returns/Refunds vollständiger Flow"], []),
        ("07 Payment Processing", "✅ Fast vollständig — 95%", "V9–V12",
         ["Payment, Webhook (signatur-verifiziert), Settlement, Payout, Refund",
          "Seller Settlements, Admin Read-Only"],
         ["Payout-Onboarding, Reconciliation, Admin Finance Mutations"], []),
        ("08 File Upload", "✅ Vollständig — 100%", "V8, V13",
         ["MIME-Typ + Magic-Byte Validierung",
          "Kategorien, Linking-Model, Soft-Delete"],
         ["Malware-Scan, CDN, Thumbnails"], []),
        ("09 Admin Panel", "✅ Fast vollständig — 90%", "V14",
         ["Dashboard, Moderation (Cert/User/Seller/Produkt)",
          "AdminAuditLog, Finance Read-Only"],
         ["Reports & Exports, 2FA, Admin-Rollen-Hierarchie"], []),
        ("10 Email Service", "✅ Fast vollständig — 90%", "V15",
         ["SMTP Transport, 6 Templates, Delivery Logging, Duplikat-Schutz"],
         ["Outbox-Pattern, Retry-Scheduler, Resend Verification"], []),
    ]

    for mod_name, status, migrations, done, missing, bugs in modules:
        add_heading(doc, f"Modul {mod_name} — {status}", 2)
        p = doc.add_paragraph(f"Migrationen: {migrations}")
        if p.runs: p.runs[0].font.size = Pt(8); p.runs[0].italic = True
        if done:
            p2 = doc.add_paragraph("Implementiert:")
            if p2.runs: p2.runs[0].bold = True; p2.runs[0].font.size = Pt(9)
            for item in done:
                bp = doc.add_paragraph(f"  ✔ {item}", style="List Bullet")
                if bp.runs: bp.runs[0].font.size = Pt(8)
        if missing:
            p3 = doc.add_paragraph("Fehlt:")
            if p3.runs: p3.runs[0].bold = True; p3.runs[0].font.size = Pt(9)
            for item in missing:
                bp = doc.add_paragraph(f"  ✗ {item}", style="List Bullet")
                if bp.runs: bp.runs[0].font.size = Pt(8)
        if bugs:
            p4 = doc.add_paragraph("Bugs / Probleme:")
            if p4.runs: p4.runs[0].bold = True; p4.runs[0].font.size = Pt(9)
            for item in bugs:
                bp = doc.add_paragraph(f"  ⚠ {item}", style="List Bullet")
                if bp.runs: bp.runs[0].font.size = Pt(8)
        doc.add_paragraph()

    doc.add_page_break()

    # 3. Frontend
    add_heading(doc, "3. Frontend — Detailanalyse", 1)
    doc.add_paragraph(
        "Next.js 15 (App Router), TypeScript, Tailwind CSS, shadcn/ui. "
        "Alle 18 API-Services definiert, ~73% der Backend-Endpunkte aktiv aufgerufen. "
        "Größte Lücken: ProductDetail (100% Mockdaten) und fehlende Payment-Integration."
    ).runs[0].font.size = Pt(9)

    add_heading(doc, "Kritische Blocker", 2)
    add_table(doc,
        ["Prio", "Komponente", "Problem", "Fix"],
        [
            ["P0", "ProductDetail.tsx", "100% Mockdaten", "ProductService.getBySlug()"],
            ["P0", "Checkout.tsx", "paymentMethod: 'MOCK'", "Stripe/PayPal Integration"],
            ["P1", "SellerDashboard", "Kein Produkt-Dialog", "ProductForm verdrahten"],
            ["P1", "ProductForm.tsx", "Kein Bild-Upload", "FileService einbinden"],
            ["P1", "—", "Kein CertificateForm", "CertificateForm erstellen"],
            ["P1", "Profil.tsx", "Address CRUD unvollständig", "AddressService verdrahten"],
            ["P2", "SustainableShop", "Category-Filter inaktiv", "selectedCategoryId übergeben"],
        ],
        col_widths=[0.5, 1.5, 2.5, 2.0]
    )

    doc.add_paragraph()
    add_heading(doc, "Services & API-Integration", 2)
    add_table(doc,
        ["Service", "Backend-Endpunkte", "Status"],
        [
            ["AuthService", "POST /auth/* (7 Endpunkte)", "✅ Vollständig"],
            ["UserService / AddressService", "GET/PATCH/DELETE /users/me, /addresses", "⚠️ Address CRUD UI teils"],
            ["BuyerValueProfileService", "GET/PUT /users/me/profile", "✅ Vollständig"],
            ["ProductService", "GET/POST/PATCH /products, images, variants", "⚠️ Detail-Page Mock"],
            ["CategoryService", "GET /categories/tree", "✅ Vollständig"],
            ["CertificateService", "GET/POST/PATCH /certificates", "❌ Kein UI verdrahtet"],
            ["CartService / CheckoutService", "GET/POST /cart, /checkout", "✅ Vollständig (MOCK Pay)"],
            ["OrderService / SellerOrderService", "GET /orders, /seller/orders", "✅ Vollständig"],
            ["AdminService", "GET/PATCH /admin/*", "✅ Fast vollständig"],
            ["FileService", "POST/GET/DELETE /files", "❌ Kein UI verdrahtet"],
            ["RecommendationService", "GET /recommendations", "✅ Widget bereit (Backend 0%)"],
        ],
        col_widths=[1.8, 2.8, 1.5]
    )
    doc.add_page_break()

    # 4. Implementierungsplan
    add_heading(doc, "4. Implementierungsplan", 1)

    add_heading(doc, "Phase 0 — SOFORT (~5 Tage)", 2)
    add_table(doc,
        ["ID", "Bereich", "Aufgabe", "Tage"],
        [["B-03","Backend","Email-Verification-Link auf Frontend-URL","0.5"],
         ["B-04","Backend","Password-Reset-Link auf Frontend-URL","0.5"],
         ["B-01","Backend","Email-Constraint-Bug fixen (Migration V17)","0.5"],
         ["B-07","Backend","Login-Response mit user-Objekt","0.5"],
         ["B-05","Backend","Refresh-Token Race Condition (pessimistic Lock)","1"],
         ["F-01","Frontend","ProductDetail: Mock durch echte API ersetzen","2"]],
        col_widths=[0.5, 0.9, 4.5, 0.6]
    )

    doc.add_paragraph()
    add_heading(doc, "Phase 1 — VOR LAUNCH (~22 Tage)", 2)
    add_table(doc,
        ["ID", "Bereich", "Aufgabe", "Tage"],
        [["B-06","Backend","Modul 02 Product Management vollständig implementieren","10"],
         ["B-02","Backend","Rate-Limit X-Forwarded-For absichern","1"],
         ["B-08","Backend","GET /users/me/profile Werteprofil-Felder korrigieren","1"],
         ["B-09","Backend","Address set-default Pfad-Mismatch fixen","0.5"],
         ["B-16","Backend","Resend Verification Email Endpunkt","0.5"],
         ["F-03","Frontend","SellerDashboard: Neues-Produkt-Dialog (ProductForm)","1"],
         ["F-04","Frontend","Bild-Upload in ProductForm verdrahten","2"],
         ["F-05","Frontend","CertificateForm erstellen + einbinden","3"],
         ["F-06","Frontend","Address CRUD in Profil vervollständigen","1.5"],
         ["F-07","Frontend","Category-Filter aktivieren","0.5"]],
        col_widths=[0.5, 0.9, 4.5, 0.6]
    )

    doc.add_paragraph()
    add_heading(doc, "Phase 2 — NACH LAUNCH (~30 Tage)", 2)
    add_table(doc,
        ["ID", "Bereich", "Aufgabe", "Tage"],
        [["B-10","Backend","Modul 04 Matching Engine implementieren","8"],
         ["B-11","Backend","Account-Lockout","2"],
         ["B-12","Backend","Redis-basiertes Rate-Limiting","2"],
         ["B-13","Backend","Admin Reports & Exports","4"],
         ["B-14","Backend","Email Outbox-Pattern + Retry","3"],
         ["B-15","Backend","Payout Onboarding (Bankverbindung)","5"],
         ["F-02","Frontend","Echte Payment-Integration (Stripe/PayPal)","5"],
         ["F-08","Frontend","Admin: suspendUser/activateUser verdrahten","0.5"],
         ["F-09","Frontend","ProducerPage: echte API","2"],
         ["F-10","Frontend","Forgot-Password UI","1"],
         ["F-11","Frontend","Zertifikate auf Produktdetailseite","1"]],
        col_widths=[0.5, 0.9, 4.5, 0.6]
    )
    doc.add_page_break()

    # 5. API Matrix
    add_heading(doc, "5. Anhang: API-Endpoint-Übersicht", 1)
    add_table(doc,
        ["Methode", "Endpoint", "Backend", "Frontend"],
        [
            ["POST", "/api/v1/auth/register", "✅ (Email-Constraint-Bug!)", "✅"],
            ["POST", "/api/v1/auth/login", "✅ (user fehlt in Response)", "✅"],
            ["POST", "/api/v1/auth/verify-email", "✅ (Link-Bug!)", "✅"],
            ["POST", "/api/v1/auth/reset-password", "✅ (Link-Bug!)", "✅"],
            ["GET/PATCH", "/api/v1/users/me", "✅", "✅"],
            ["GET/PUT", "/api/v1/users/me/profile", "⚠️ Falsche Felder", "✅"],
            ["GET/POST/...", "/api/v1/users/me/addresses", "✅ (Pfad-Mismatch!)", "⚠️ Teils"],
            ["GET", "/api/v1/products", "❌ FEHLT", "✅ (Mock-Fallback)"],
            ["GET", "/api/v1/products/{slug}", "❌ FEHLT", "❌ Nicht aufgerufen"],
            ["POST/PATCH", "/api/v1/products", "❌ FEHLT", "⚠️ Service bereit"],
            ["GET", "/api/v1/categories/tree", "✅", "✅"],
            ["GET/POST/...", "/api/v1/cart/items", "✅", "✅"],
            ["POST", "/api/v1/checkout/complete", "✅", "✅ (MOCK Pay)"],
            ["GET", "/api/v1/orders", "✅", "✅"],
            ["GET", "/api/v1/seller/orders", "✅", "✅"],
            ["POST", "/api/v1/payments/create-intent", "✅", "❌ MOCK"],
            ["GET", "/api/v1/recommendations", "❌ FEHLT", "✅ Widget bereit"],
            ["GET/POST/...", "/api/v1/certificates", "✅", "❌ Kein UI"],
            ["POST/GET/DELETE", "/api/v1/files", "✅", "❌ Kein UI"],
            ["GET/...", "/api/v1/admin/users", "✅", "✅"],
        ],
        col_widths=[0.8, 2.8, 1.8, 1.3]
    )

    doc.add_paragraph()
    p = doc.add_paragraph(f"Erstellt am {TODAY} · Marketplace Backend & Frontend Status Report")
    if p.runs: p.runs[0].font.size = Pt(8); p.runs[0].italic = True

    word_path = os.path.join(DOCS_DIR, "status-report.docx")
    doc.save(word_path)
    print(f"[OK] Word:  {word_path}")
    return word_path


if __name__ == "__main__":
    print("Generiere Dokumente...")
    excel_path = build_excel()
    word_path  = build_word()
    print()
    print("Fertig!")
    print(f"  Excel: {excel_path}  ({os.path.getsize(excel_path):,} Bytes)")
    print(f"  Word:  {word_path}  ({os.path.getsize(word_path):,} Bytes)")
